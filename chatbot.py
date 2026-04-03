import os
import re
import math
import base64
import logging
import json
import uvicorn
from typing import List, Dict, Optional
from io import BytesIO

# Third-party libraries
from dotenv import load_dotenv
from openai import OpenAI
from pypdf import PdfReader
from docx import Document
from PIL import Image

# Setup simple logging
logging.basicConfig(level=logging.INFO, format='%(levelname)s: %(message)s')
logger = logging.getLogger(__name__)

def load_system_message_from_json(file_path: str) -> str:
    """Helper to load a custom system prompt from a JSON file."""
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)
            return data.get("system_message", "")
    except Exception:
        return "You are a helpful university assistant."

FALLBACK_MODEL = "llama-3.1-8b-instant"  # Lighter model when rate-limited

class Chatbot:
    def __init__(self, model: str = "llama-3.3-70b-versatile", vision_model: str = "meta-llama/llama-4-scout-17b-16e-instruct", data_dir: str = "data"):
        load_dotenv()
        self.api_key = os.getenv("GROQ_API_KEY", "missing_key")
        if self.api_key == "missing_key":
            logger.warning("GROQ_API_KEY not found in .env file.")
        
        # We use a separate client/model for Vision specifically
        self.client = OpenAI(
            base_url="https://api.groq.com/openai/v1",
            api_key=self.api_key
        )
        self.model = model
        self.vision_model = vision_model
        self.store: List[Dict] = []
        self.data_dir = data_dir
        
        if not os.path.exists(self.data_dir):
            os.makedirs(self.data_dir)

    def load_document_from_bytes(self, filename: str, file_bytes: bytes):
        """Load a document from raw bytes — used for Supabase Storage downloads."""
        try:
            text = self._read_bytes(filename, file_bytes)
            if text.strip():
                self._chunk_and_store(text, filename, is_official=False)
        except Exception as e:
            logger.error(f"Error loading {filename} from bytes: {e}")

    def _read_bytes(self, filename: str, file_bytes: bytes) -> str:
        """Parse document content from raw bytes."""
        ext = os.path.splitext(filename)[1].lower()
        if ext in [".txt", ".md"]:
            return file_bytes.decode("utf-8", errors="ignore")
        elif ext == ".pdf":
            reader = PdfReader(BytesIO(file_bytes))
            return "\n".join([p.extract_text() for p in reader.pages if p.extract_text()])
        elif ext == ".docx":
            doc = Document(BytesIO(file_bytes))
            return "\n".join([p.text for p in doc.paragraphs])
        elif ext in [".jpg", ".jpeg", ".png"]:
            return self._ocr_vision_bytes(file_bytes, filename)
        return ""

    def _ocr_vision_bytes(self, file_bytes: bytes, filename: str) -> str:
        """Run Vision OCR on image bytes."""
        logger.info(f"Performing Vision OCR on: {filename}")
        try:
            with Image.open(BytesIO(file_bytes)) as img:
                if img.mode in ("RGBA", "P"):
                    img = img.convert("RGB")
                buffered = BytesIO()
                img.save(buffered, format="JPEG", quality=85)
                b64 = base64.b64encode(buffered.getvalue()).decode("utf-8")
            response = self.client.chat.completions.create(
                model=self.vision_model,
                messages=[{"role": "user", "content": [
                    {"type": "text", "text": "EXTRACT all written text from this image exactly as it appears. Output ONLY the transcribed text."},
                    {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{b64}"}}
                ]}],
                max_tokens=1024,
            )
            transcription = response.choices[0].message.content
            logger.info(f"OCR Complete: {len(transcription)} chars.")
            return f"--- [OCR Content from {filename}] ---\n{transcription}"
        except Exception as e:
            logger.error(f"Vision OCR failed: {e}")
            return f"[Error: Could not read image {filename}]"

    def load_documents(self):
        """Initial data ingestion from the data folder. 
        Differentiates between 'official' university files and 'user' study notes.
        """
        if not os.path.exists(self.data_dir):
            return
            
        self.store = []
        # Define which files are 'official' university admin files
        OFFICIAL_DOCS = ["about_university.txt", "academic_calendar.txt", "contact_info.txt", "system_message.json"]
        
        files = [f for f in os.listdir(self.data_dir) if os.path.isfile(os.path.join(self.data_dir, f))]
        
        for filename in files:
            file_path = os.path.join(self.data_dir, filename)
            is_official = filename in OFFICIAL_DOCS
            try:
                text = self._read_file(file_path)
                if text.strip():
                    self._chunk_and_store(text, filename, is_official=is_official)
                    type_str = "OFFICIAL" if is_official else "STUDY NOTE"
                    logger.info(f"Loaded {type_str}: {filename}")
            except Exception as e:
                logger.error(f"Error loading {filename}: {e}")
        
        logger.info(f"Knowledge Base Synced: {len(self.store)} chunks total.")

    def _read_file(self, file_path: str) -> str:
        """Reads multiple formats: TXT, MD, PDF, DOCX, and now JPG/PNG (OCR via Vision)."""
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext in [".txt", ".md"]:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        
        elif ext == ".pdf":
            reader = PdfReader(file_path)
            return "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
            
        elif ext == ".docx":
            doc = Document(file_path)
            return "\n".join([p.text for p in doc.paragraphs])
            
        elif ext in [".jpg", ".jpeg", ".png"]:
            return self._ocr_vision(file_path)
            
        return ""

    def _ocr_vision(self, file_path: str) -> str:
        """Uses Groq's Vision model to extract text from handwriting/photos."""
        logger.info(f"Performing Vision OCR on: {os.path.basename(file_path)}")
        try:
            # 1. Prepare Image (resizing slightly can help with token limits)
            with Image.open(file_path) as img:
                # Convert to RGB if needed (removes alpha channel)
                if img.mode in ("RGBA", "P"):
                    img = img.convert("RGB")
                
                # Buffer for base64
                buffered = BytesIO()
                img.save(buffered, format="JPEG", quality=85)
                base64_image = base64.b64encode(buffered.getvalue()).decode('utf-8')

            # 2. Call Groq Vision API
            response = self.client.chat.completions.create(
                model=self.vision_model,
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": "EXTRACT all written text from this image exactly as it appears. If it's a student note, preserve the structure (bullets, headers). Output ONLY the transcribed text."},
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/jpeg;base64,{base64_image}",
                                },
                            },
                        ],
                    }
                ],
                max_tokens=1024,
            )
            transcription = response.choices[0].message.content
            logger.info(f"OCR Complete: Extracted {len(transcription)} chars.")
            return f"--- [OCR Content from {os.path.basename(file_path)}] ---\n{transcription}"
            
        except Exception as e:
            logger.error(f"Vision OCR failed: {e}")
            return f"[Error: Could not read image {os.path.basename(file_path)}]"

    def _chunk_and_store(self, text: str, source: str, is_official: bool = False, chunk_size: int = 800):
        """Simple overlap chunking with tags."""
        for i in range(0, len(text), chunk_size - 150):
            chunk = text[i : i + chunk_size]
            self.store.append({
                "content": chunk, 
                "source": source, 
                "is_official": is_official
            })

    def _retrieve_context(self, query: str, top_k: int = 6) -> List[Dict]:
        """Keyword-based ranking — searches ALL documents (official + user uploads)."""
        if not self.store: return []

        query_words = set(re.findall(r'\w+', query.lower()))
        scored = []
        for doc in self.store:
            content_lower = doc["content"].lower()
            score = sum(1 for word in query_words if word in content_lower)
            # Boost official docs slightly so KB knowledge surfaces
            if doc.get("is_official", False):
                score += 1
            if score > 0:
                scored.append((score, doc))

        scored.sort(key=lambda x: x[0], reverse=True)
        return [item[1] for item in scored[:top_k]]

    def get_response(self, user_query: str, mode: str = "chat"):
        """Unified chat — searches all available knowledge and responds helpfully."""
        context_docs = self._retrieve_context(user_query)

        # Fallback: if no keyword match but docs exist, pull recent ones
        if not context_docs and self.store:
            context_docs = self.store[-4:]

        context_str = "\n---\n".join([f"Source: {d['source']}\n{d['content']}" for d in context_docs])
        sources = list(set([d["source"] for d in context_docs]))

        system_msg = (
            "You are Academix, a friendly and knowledgeable university learning assistant. "
            "Your purpose is to help students learn, understand concepts, and prepare for exams.\n\n"
            "RULES:\n"
            "- ALWAYS try to help the student with their question using your own knowledge.\n"
            "- If relevant context from uploaded notes or the knowledge base is provided below, use it to give more specific answers.\n"
            "- If NO context is provided, still help the student using your general knowledge. NEVER refuse to help or repeatedly ask them to upload files.\n"
            "- Be conversational, clear, and structured. Use bullet points, headings, and examples.\n"
            "- If the student asks about exam preparation, help them with study strategies, key concepts, and practice questions.\n"
            "- Keep answers focused and practical — students want to learn, not read walls of text.\n"
        )

        if context_str.strip():
            system_msg += f"\nAVAILABLE KNOWLEDGE CONTEXT:\n{context_str}"
        else:
            system_msg += "\nNote: No documents have been uploaded yet. Answer using your general knowledge. You may suggest the student upload their notes for more personalized help, but do NOT make it a requirement — help them regardless."

        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": system_msg},
                    {"role": "user",   "content": user_query}
                ],
                temperature=0.6
            )
            return {
                "answer":  response.choices[0].message.content,
                "sources": sources
            }
        except Exception as e:
            error_str = str(e).lower()
            # Rate limit / token quota exceeded → retry with lighter model
            if "rate_limit" in error_str or "rate limit" in error_str or "token" in error_str or "429" in error_str or "quota" in error_str:
                logger.warning(f"Rate limited on {self.model}, falling back to {FALLBACK_MODEL}")
                try:
                    response = self.client.chat.completions.create(
                        model=FALLBACK_MODEL,
                        messages=[
                            {"role": "system", "content": system_msg},
                            {"role": "user",   "content": user_query}
                        ],
                        temperature=0.6
                    )
                    return {
                        "answer": response.choices[0].message.content + "\n\n> *Switched to a faster model due to high demand. Responses may be shorter.*",
                        "sources": sources
                    }
                except Exception as fallback_err:
                    logger.error(f"Fallback model also failed: {fallback_err}")
            logger.error(f"Chat API Error: {e}")
            return {"answer": f"Error: {str(e)}", "sources": []}

    def vision_chat(self, user_query: str, image_bytes_list: list, mode: str = "chat"):
        """Send images directly to the vision model with the user's question.
        This gives the AI real eyes — it sees the image and answers about it.
        """
        # Build the user content array with text + images
        user_content = [{"type": "text", "text": user_query}]
        
        for img_bytes in image_bytes_list:
            b64 = base64.b64encode(img_bytes).decode("utf-8")
            user_content.append({
                "type": "image_url",
                "image_url": {"url": f"data:image/jpeg;base64,{b64}"}
            })
        
        system_msg = (
            "You are Academix, a smart university assistant. The student has attached image(s) to their message. "
            "Analyze the image(s) carefully and respond to their question. "
            "If the image contains notes, diagrams, or text — read and explain it. "
            "If it's a screenshot of a problem — solve it. Be helpful, structured, and academic."
        )
        
        try:
            response = self.client.chat.completions.create(
                model=self.vision_model,
                messages=[
                    {"role": "system", "content": system_msg},
                    {"role": "user", "content": user_content}
                ],
                max_tokens=2048,
                temperature=0.5
            )
            return {
                "answer": response.choices[0].message.content,
                "sources": ["📷 Attached Image(s)"]
            }
        except Exception as e:
            error_str = str(e)
            logger.error(f"Vision Chat Error: {error_str}")

            # Rate limit on vision → fall back to text-only with lighter model
            error_lower = error_str.lower()
            if "rate_limit" in error_lower or "rate limit" in error_lower or "429" in error_lower or "quota" in error_lower:
                logger.warning(f"Vision rate limited, falling back to {FALLBACK_MODEL} text-only")
                try:
                    response = self.client.chat.completions.create(
                        model=FALLBACK_MODEL,
                        messages=[
                            {"role": "system", "content": f"The user attached {len(image_bytes_list)} image(s) but the system is rate-limited. Help with their text query. Mention you couldn't process the image due to high demand."},
                            {"role": "user", "content": user_query}
                        ],
                        temperature=0.7
                    )
                    return {
                        "answer": response.choices[0].message.content + "\n\n> *High demand — image analysis temporarily unavailable. Try again in a moment.*",
                        "sources": ["📷 Image received (rate limited)"]
                    }
                except:
                    pass

            # If the vision model is blocked/unavailable, fall back to text-only with a helpful message
            if "permission" in error_lower or "decommissioned" in error_lower or "not found" in error_lower:
                # Fall back to text model acknowledging the image
                fallback_msg = (
                    f"The user attached {len(image_bytes_list)} image(s) with their message but the vision model is currently unavailable. "
                    f"Acknowledge that you received image attachments but explain you cannot view them right now. "
                    f"Still try to help with their text query. Suggest they describe the image content to you instead."
                )
                try:
                    response = self.client.chat.completions.create(
                        model=self.model,
                        messages=[
                            {"role": "system", "content": fallback_msg},
                            {"role": "user", "content": user_query}
                        ],
                        temperature=0.7
                    )
                    return {
                        "answer": response.choices[0].message.content + "\n\n> ⚙️ *Vision model is currently unavailable. To enable image analysis, go to [console.groq.com/settings/limits](https://console.groq.com/settings/limits) and enable the Llama 4 Scout model.*",
                        "sources": ["📷 Image received (vision offline)"]
                    }
                except:
                    pass
            
            return {"answer": f"Error processing image: {error_str}", "sources": []}

    def chat(self, query: str, mode: str = "chat", image_bytes_list: list = None):
        """Main entry point. Routes to vision model if images are attached."""
        if image_bytes_list:
            return self.vision_chat(query, image_bytes_list, mode=mode)
        return self.get_response(query, mode=mode)
